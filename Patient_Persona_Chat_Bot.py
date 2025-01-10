import openai
import streamlit as st
from openai import AzureOpenAI
# from dotenv import load_dotenv
import os
from openai import OpenAI
from io import BytesIO
import base64
# from audio_recorder_streamlit import audio_recorder
from st_audiorec import st_audiorec
from streamlit_float import *
from supabase import create_client, Client
import re
from openai.types.beta.assistant_stream_event import ThreadMessageDelta
from openai.types.beta.threads.text_delta_block import TextDeltaBlock
from docx import Document
import toml
import warnings

warnings.filterwarnings("ignore", category=DeprecationWarning)

def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def speech_to_text(client_speech, audio_data):
    with open(audio_data, "rb") as audio_file:
        transcript = client_speech.audio.transcriptions.create(
            model="whisper-1",
            response_format="text",
            file=audio_file
        )
    return transcript

def text_to_speech(client_speech, input_text):
    response = client_speech.audio.speech.create(
        model="tts-1",
        voice="nova",
        input=input_text
    )
    webm_file_path = "temp_audio_play.mp3"
    with open(webm_file_path, "wb") as f:
        response.stream_to_file(webm_file_path)
    return webm_file_path

def autoplay_audio(file_path: str):
    with open(file_path, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode("utf-8")
    md = f"""
    <audio autoplay>
    <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
    </audio>
    """
    st.markdown(md, unsafe_allow_html=True)

# Function to generate Word document from the transcript
def create_transcript_word(transcript):
    doc = Document()
    doc.add_heading("Nurse - Patient Conversation Transcript\n", level=1)

    for line in transcript:
        if line.startswith("Patient:"):
            p = doc.add_paragraph()
            p.add_run("Patient:").bold = True
            p.add_run(line[len("Patient:"):])
        elif line.startswith("Student Nurse:"):
            p = doc.add_paragraph()
            p.add_run("Student Nurse:").bold = True
            p.add_run(line[len("Student Nurse:"):])
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to initialize or retrieve the assistant and thread
def init_thread(client_azure):
    # Create a Thread
    st.session_state.thread = client_azure.beta.threads.create()
    # st.session_state.messages = []

def main():
    # load_dotenv('patient_persona.env')
    # Inject CSS for custom styles
    local_css("patient_persona_style.css")

    SUPABASE_URL = st.secrets['SUPABASE_URL']
    SUPABASE_KEY = st.secrets['SUPABASE_KEY']
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

    st.session_state.assistant_id = "asst_3kMmf6f7HKR8i14CJSZy9uPf"

    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'full_transcript' not in st.session_state:
        st.session_state.full_transcript = []
    if 'processed_audio' not in st.session_state:
        st.session_state.processed_audio = None  # Initialize the processed audio state
    if 'assistant_instruction' not in st.session_state:
        st.session_state.assistant_instruction = None
    if 'assistant_id' not in st.session_state:
        st.session_state.assistant_id = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "manual_input" not in st.session_state:
        st.session_state.manual_input = None  # Add this line to initialize manual input
    if "end_session_button_clicked" not in st.session_state:
        st.session_state.end_session_button_clicked = False  # Initialize the state to track the button click

    st.title("AI SimBot")

    client_azure = AzureOpenAI(
        azure_endpoint=st.secrets["AZURE_ENDPOINT"],
        api_key=st.secrets["AZURE_OPENAI_KEY"],
        api_version="2024-02-15-preview"
    )

    client_speech = OpenAI(api_key=st.secrets['OPENAI_API_KEY'])

    response = supabase.table("Patient_Persona_Instructions").select("*").eq('Patient_Persona_Active_Flag',
                                                                             'Y').execute()
    data = response.data

    if data and isinstance(data, list) and len(data) > 0:
        patient_persona_name = data[0].get("Patient_Persona_Name")
        patient_persona_age = data[0].get("Patient_Persona_Age")
        patient_persona_gender = data[0].get("Patient_Persona_Gender")
        patient_persona_visit_reason = data[0].get("Patient_Persona_Visit_Reason")
        patient_persona_background_information = data[0].get('Patient_Persona_Background_Information')
        patient_persona_manual_input = data[0].get('Patient_Persona_Manual_Input')
        patient_persona_ai_generation = data[0].get('Patient_Persona_AI_Generation')
        patient_persona_introduction_text = data[0].get('Patient_Persona_Introduction_Text')

    # if selected_tab == "Manual Input":
    st.sidebar.header("Chat with your Patient")
    container1 = st.sidebar.container(border=True)
    with container1:
        st.subheader(f"Name: {patient_persona_name}")
        st.subheader(f"Age: {patient_persona_age}")
        st.subheader(f"Gender: {patient_persona_gender}")
        st.subheader(f"Reason for Visit: {patient_persona_visit_reason}")
        st.image('Patient_Persona_Assets/Jordan_Avatar.jpg')

    if patient_persona_manual_input:
        # Perform replacements for each placeholder
        patient_persona_manual_input_modified = re.sub(r'\{patient_persona_name\}', patient_persona_name,
                                                       patient_persona_manual_input)
        patient_persona_manual_input_modified = re.sub(r'\{patient_persona_age\}', str(patient_persona_age),
                                                       patient_persona_manual_input_modified)
        patient_persona_manual_input_modified = re.sub(r'\{patient_persona_gender\}', patient_persona_gender,
                                                       patient_persona_manual_input_modified)
        patient_persona_manual_input_modified = re.sub(r'\{patient_persona_visit_reason\}',
                                                       patient_persona_visit_reason,
                                                       patient_persona_manual_input_modified)
        patient_persona_manual_input_modified = re.sub(r'\{patient_persona_background_information\}',
                                                       patient_persona_background_information,
                                                       patient_persona_manual_input_modified)

        init_thread(client_azure)
        st.session_state.assistant_instruction = patient_persona_manual_input_modified

    # Button to start the chat session
    col1, col2, col3 = st.sidebar.columns([1, 1, 1])
    with col2:
        if st.button('Start Chat'):
            st.session_state["chat_active"] = True
            st.session_state.messages = []
            st.session_state.conversation_history = []
            st.session_state.full_transcript = []
            st.session_state.chat_history = []
            st.session_state.show_text = False  # Hide the text when Start Chat is pressed

    # Display text before Start Chat is pressed
    if "show_text" not in st.session_state:
        st.session_state.show_text = True

    if st.session_state.show_text:
        container2 = st.container(border=True)
        with container2:
            st.markdown(patient_persona_introduction_text)

    # Check if chat is active
    if st.session_state.get("chat_active"):
        st.sidebar.header('Speech Input')

        for message in st.session_state.chat_history:
            if message["role"] == 'user':
                with st.chat_message(message["role"], avatar="Patient_Persona_Assets/Nurse_Avatar.png"):
                    st.markdown(message["content"])
            else:
                with st.chat_message(message["role"], avatar="Patient_Persona_Assets/User_Avatar.jpg"):
                    st.markdown(message["content"])

        # Check if there's a manual input and process it
        if st.session_state.manual_input:
            user_query = st.session_state.manual_input
            st.session_state.manual_input = None  # Clear manual input after use
        else:
            if st.session_state.end_session_button_clicked:
                user_query = None  # Prevent any input after the end session
            else:
                user_query = st.chat_input("Click 'End Session' Button to Receive Feedback and Download Transcript.")

        # Create footer container for the microphone
        footer_container = st.sidebar.container(border=True)

        with footer_container:
            # audio_bytes = audio_recorder(text="Click to Record", icon_size="2x", neutral_color="#30475E")
            audio_bytes = st_audiorec()
            # st.write(audio_bytes)

            # Check if there is a new audio recording and it has not been processed yet
            if audio_bytes and audio_bytes != st.session_state.processed_audio and audio_bytes != b'RIFF,\x00\x00\x00WAVEfmt \x10\x00\x00\x00\x01\x00\x02\x00\x80\xbb\x00\x00\x00\xee\x02\x00\x04\x00\x10\x00data\x00\x00\x00\x00':
                webm_file_path = "temp_audio.mp3"
                with open(webm_file_path, "wb") as f:
                    f.write(audio_bytes)
                transcript = speech_to_text(client_speech, webm_file_path)
                if transcript:
                    os.remove(webm_file_path)
                    user_query = transcript  # Ensure you extract the text from the transcript object correctly
                    st.session_state.processed_audio = audio_bytes  # Update the processed audio state

            if audio_bytes == b'RIFF,\x00\x00\x00WAVEfmt \x10\x00\x00\x00\x01\x00\x02\x00\x80\xbb\x00\x00\x00\xee\x02\x00\x04\x00\x10\x00data\x00\x00\x00\x00':
                st.error('Something went wrong. Kindly refresh the page and try again.')
                st.stop()

        if user_query:
            if user_query.lower() == 'exit':
                st.session_state["chat_active"] = False
                st.stop()

            # Create a new thread if it does not exist
            if "thread_id" not in st.session_state:
                thread = client_azure.beta.threads.create()
                st.session_state.thread_id = thread.id

            # Display the user's query
            with st.chat_message("user", avatar="Patient_Persona_Assets/Nurse_Avatar.png"):
                st.markdown(user_query)

            # Store the user's query into the history
            st.session_state.chat_history.append({"role": "user", "content": user_query})
            role = 'Student Nurse'
            st.session_state.full_transcript.append(f"{role}: {user_query}")

            # Add user query to the thread
            client_azure.beta.threads.messages.create(
                thread_id=st.session_state.thread_id,
                role="user",
                content=user_query
            )

            # Stream the assistant's reply
            with st.chat_message("assistant", avatar="Patient_Persona_Assets/User_Avatar.jpg"):
                stream = client_azure.beta.threads.runs.create(
                    thread_id=st.session_state.thread_id,
                    assistant_id=st.session_state.assistant_id,
                    instructions=st.session_state.assistant_instruction,
                    stream=True,
                )

                # Empty container to display the assistant's reply
                assistant_reply_box = st.empty()

                # A blank string to store the assistant's reply
                assistant_reply = ""

                # Iterate through the stream
                for event in stream:
                    if isinstance(event, ThreadMessageDelta):
                        if isinstance(event.data.delta.content[0], TextDeltaBlock):
                            # empty the container
                            assistant_reply_box.empty()
                            # add the new text
                            assistant_reply += event.data.delta.content[0].text.value
                            # display the new text
                            assistant_reply_box.markdown(assistant_reply)
                    if event.data.object == "thread.message.delta":
                        for content in event.data.delta.content:
                            if content.type == 'image_file':
                                file_id = content.image_file.file_id
                                image_data = client_azure.files.content(file_id)
                                image_data_bytes = image_data.read()
                                with open("my-image.png", "wb") as file:
                                    file.write(image_data_bytes)
                                st.image("my-image.png")

                # Once the stream is over, update chat history
                st.session_state.chat_history.append({"role": "assistant", "content": assistant_reply})
                role = 'Patient'
                st.session_state.full_transcript.append(f"{role}: {assistant_reply}")

                audio_file = text_to_speech(client_speech, assistant_reply)
                autoplay_audio(audio_file)



    st.sidebar.warning('Please be advised that this patient simulation includes themes and discussions of substance use. This simulation might evoke strong emotions or may be distressing for individuals with personal experiences related to substance use or addiction. If you feel overwhelmed by the content, please discontinue the simulation and reach out to your course faculty, or schedule an appointment with [Northeastern University Health & Counseling Services](https://uhcs.northeastern.edu/). If you or someone you know is struggling with substance use or addiction, support is available.')

    # Button to show the download button
    if st.session_state.get("chat_active"):
        if st.session_state.end_session_button_clicked:
            st.button("End Session", disabled=True)
        else:
            if st.session_state.get("full_transcript"):
                if st.button("End Session"):
                    st.session_state.end_session_button_clicked = True
                    st.session_state["end_session_button"] = True
                    st.session_state.manual_input = "Goodbye. Thank you for coming."  # Set manual input
                    # Trigger the manual input immediately
                    user_query = st.session_state.manual_input
                    st.session_state.manual_input = None  # Clear manual input after use
                    # Process the manual input as a regular chat input
                    if "thread_id" not in st.session_state:
                        thread = client_azure.beta.threads.create()
                        st.session_state.thread_id = thread.id

                    with st.chat_message("user", avatar="Patient_Persona_Assets/Nurse_Avatar.png"):
                        st.markdown(user_query)

                    st.session_state.chat_history.append({"role": "user", "content": user_query})
                    role = 'Student Nurse'
                    st.session_state.full_transcript.append(f"{role}: {user_query}")

                    client_azure.beta.threads.messages.create(
                        thread_id=st.session_state.thread_id,
                        role="user",
                        content=user_query
                    )

                    with st.chat_message("assistant", avatar="Patient_Persona_Assets/User_Avatar.jpg"):
                        stream = client_azure.beta.threads.runs.create(
                            thread_id=st.session_state.thread_id,
                            assistant_id=st.session_state.assistant_id,
                            instructions=st.session_state.assistant_instruction,
                            stream=True,
                        )

                        assistant_reply_box = st.empty()
                        assistant_reply = ""

                        for event in stream:
                            if isinstance(event, ThreadMessageDelta):
                                if isinstance(event.data.delta.content[0], TextDeltaBlock):
                                    assistant_reply_box.empty()
                                    assistant_reply += event.data.delta.content[0].text.value
                                    assistant_reply_box.markdown(assistant_reply)
                            if event.data.object == "thread.message.delta":
                                for content in event.data.delta.content:
                                    if content.type == 'image_file':
                                        file_id = content.image_file.file_id
                                        image_data = client_azure.files.content(file_id)
                                        image_data_bytes = image_data.read()
                                        with open("my-image.png", "wb") as file:
                                            file.write(image_data_bytes)
                                        st.image("my-image.png")

                        st.session_state.chat_history.append({"role": "assistant", "content": assistant_reply})
                        role = 'Patient'
                        st.session_state.full_transcript.append(f"{role}: {assistant_reply}")

    # Check if there's a transcript to download
    if st.session_state.get("full_transcript") and st.session_state.get("end_session_button"):
        full_transcript_text = "\n".join(st.session_state.full_transcript)
        word_buffer = create_transcript_word(st.session_state.full_transcript)

        col1, col2 = st.columns([1, 1])
        # Button to download the full conversation transcript
        with col1:
            st.download_button(
                label="Download Transcript",
                data=word_buffer,
                file_name="Nurse_Patient_Full_Transcript.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
